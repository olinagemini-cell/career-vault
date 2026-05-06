"""
将庭审策略文档（律师版）.md 转换为格式良好的 Word 文档。
使用 python-docx 构建，支持标题层级、表格、引用块、代码块、加粗、列表等。
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT = Path(r"C:\Users\olina\Desktop\诉讼\庭审策略文档（律师版）.md")
OUTPUT = INPUT.with_suffix(".docx")


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_text(paragraph, text):
    """Parse inline markdown (bold, inline code) and add runs to paragraph."""
    # Split by bold markers **...**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif '`' in part:
            # Handle inline code
            code_parts = re.split(r'(`[^`]+`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = paragraph.add_run(cp[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
                else:
                    if cp:
                        paragraph.add_run(cp)
        else:
            if part:
                paragraph.add_run(part)


def add_quote_block(doc, lines):
    """Add a blockquote as an indented paragraph with left border styling."""
    text = '\n'.join(line.lstrip('> ').rstrip() for line in lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add a subtle visual indicator
    run = p.add_run('『')
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(9)
    add_formatted_text(p, text)
    run = p.add_run('』')
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(9)


def add_code_block(doc, lines):
    """Add a code/diagram block as monospaced text."""
    text = '\n'.join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_table(doc, header_line, rows):
    """Add a markdown table to the document."""
    headers = [c.strip() for c in header_line.strip('|').split('|')]
    num_cols = len(headers)

    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        add_formatted_text(p, h)
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for row_line in rows:
        cells_text = [c.strip() for c in row_line.strip('|').split('|')]
        row = table.add_row()
        for i in range(min(num_cols, len(cells_text))):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            add_formatted_text(p, cells_text[i])
            for run in p.runs:
                run.font.size = Pt(9)

    # Set column widths roughly
    doc.add_paragraph()  # spacing after table


def add_list_item(doc, text, level=0):
    """Add a list item (bullet or checkbox)."""
    # Handle checkboxes
    if text.startswith('[ ] '):
        prefix = '☐ '
        text = text[4:]
    elif text.startswith('[x] '):
        prefix = '☑ '
        text = text[4:]
    else:
        prefix = ''

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

    if prefix:
        run = p.add_run(prefix)
        run.font.size = Pt(10)
    else:
        run = p.add_run('• ')
        run.font.size = Pt(10)

    add_formatted_text(p, text)


def convert():
    md_text = INPUT.read_text(encoding='utf-8')
    lines = md_text.split('\n')

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Set heading styles
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = '微软雅黑'
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if level == 1:
            hs.font.size = Pt(18)
            hs.font.color.rgb = RGBColor(0x1F, 0x37, 0x63)
        elif level == 2:
            hs.font.size = Pt(15)
            hs.font.color.rgb = RGBColor(0x2E, 0x4A, 0x7A)
        elif level == 3:
            hs.font.size = Pt(13)
            hs.font.color.rgb = RGBColor(0x3B, 0x5E, 0x91)
        elif level == 4:
            hs.font.size = Pt(11.5)
            hs.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            # Add a subtle separator
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Headings
        if stripped.startswith('#'):
            match = re.match(r'^(#{1,4})\s+(.*)', stripped)
            if match:
                level = len(match.group(1))
                heading_text = match.group(2)
                # Remove any remaining markdown bold from heading
                heading_text = heading_text.replace('**', '')
                doc.add_heading(heading_text, level=level)
                i += 1
                continue

        # Code block
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1  # skip closing ```
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and (lines[i].strip().startswith('>') or (lines[i].strip() and quote_lines and not lines[i].strip().startswith('#') and not lines[i].strip().startswith('|') and not lines[i].strip().startswith('-'))):
                if lines[i].strip().startswith('>'):
                    quote_lines.append(lines[i])
                elif lines[i].strip():
                    quote_lines.append(lines[i])
                else:
                    break
                i += 1
            add_quote_block(doc, quote_lines)
            continue

        # Table
        if stripped.startswith('|') and '|' in stripped[1:]:
            header_line = stripped
            i += 1
            # Skip separator line (|---|---|)
            if i < len(lines) and re.match(r'^\|[\s\-:|]+\|', lines[i].strip()):
                i += 1
            # Collect data rows
            data_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                data_rows.append(lines[i].strip())
                i += 1
            add_table(doc, header_line, data_rows)
            continue

        # List items
        if re.match(r'^[-*]\s', stripped) or re.match(r'^\d+\.\s', stripped):
            # Determine text after marker
            list_match = re.match(r'^[-*]\s+(.*)', stripped)
            if not list_match:
                list_match = re.match(r'^\d+\.\s+(.*)', stripped)
            if list_match:
                list_text = list_match.group(1)
                # Determine indent level from original line
                leading_spaces = len(line) - len(line.lstrip())
                level = leading_spaces // 2
                add_list_item(doc, list_text, level)
                i += 1
                continue

        # Bold-start metadata lines (like **文档性质**: ...)
        if stripped.startswith('**') and '**' in stripped[2:]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, stripped)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.5)

    doc.save(str(OUTPUT))
    print(f"已保存: {OUTPUT}")


if __name__ == '__main__':
    convert()
