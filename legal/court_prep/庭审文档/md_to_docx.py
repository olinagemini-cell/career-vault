"""
Markdown → Word (.docx) 转换脚本
用于将庭审策略文档和发问提纲的Markdown文件转为格式化的Word文档
"""
import re, sys, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def setup_styles(doc):
    """设置中文友好的文档样式"""
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for i in range(1, 5):
        hstyle = doc.styles[f'Heading {i}']
        hstyle.font.name = '黑体'
        hstyle._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        hstyle.font.color.rgb = RGBColor(0, 0, 0)
        sizes = {1: 18, 2: 15, 3: 13, 4: 12}
        hstyle.font.size = Pt(sizes.get(i, 12))
        hstyle.paragraph_format.space_before = Pt(12 if i <= 2 else 8)
        hstyle.paragraph_format.space_after = Pt(4)

    # 引用样式
    if 'Quote' not in [s.name for s in doc.styles]:
        quote_style = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote_style = doc.styles['Quote']
    quote_style.font.name = '楷体'
    quote_style.font.size = Pt(11)
    quote_style.font.italic = True
    quote_style._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    quote_style.paragraph_format.left_indent = Cm(1)
    quote_style.paragraph_format.space_after = Pt(4)

    # 代码块样式
    code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9.5)
    code_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    code_style.paragraph_format.left_indent = Cm(0.5)
    code_style.paragraph_format.space_after = Pt(1)
    code_style.paragraph_format.line_spacing = 1.1
    # 灰色底纹
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
    code_style._element.pPr.append(shading)

    return doc


def add_rich_text(paragraph, text):
    """解析行内markdown格式（粗体、行内代码、链接等）并添加到段落"""
    # 处理粗体、行内代码
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            # 处理可能的emoji标记（保留原样）
            paragraph.add_run(part)


def add_table(doc, header_line, rows):
    """添加Word表格"""
    # 解析表头
    headers = [c.strip() for c in header_line.strip('|').split('|')]
    num_cols = len(headers)

    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        add_rich_text(p, h)
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
        # 表头灰底
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E8E8" w:val="clear"/>')
        cell._element.tcPr.append(shading)

    # 数据行
    for row_text in rows:
        cols = [c.strip() for c in row_text.strip('|').split('|')]
        row = table.add_row()
        for i in range(min(len(cols), num_cols)):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            add_rich_text(p, cols[i])
            for run in p.runs:
                run.font.size = Pt(10)

    # 设置字号
    doc.add_paragraph('')  # 表后空行
    return table


def convert_md_to_docx(md_path, docx_path):
    """主转换函数"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    setup_styles(doc)

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    i = 0
    in_code_block = False
    in_table = False
    table_header = None
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # 代码块
        if line.strip().startswith('```'):
            if in_code_block:
                in_code_block = False
                doc.add_paragraph('')  # 代码块后空行
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph(line, style='CodeBlock')
            i += 1
            continue

        # 分隔线
        if re.match(r'^---+\s*$', line):
            i += 1
            continue

        # 空行
        if not line.strip():
            # 如果在表格中，结束表格
            if in_table:
                add_table(doc, table_header, table_rows)
                in_table = False
                table_header = None
                table_rows = []
            i += 1
            continue

        # 表格检测
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                # 检查下一行是否是分隔行
                if i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i+1].strip()):
                    in_table = True
                    table_header = line
                    table_rows = []
                    i += 2  # 跳过分隔行
                    continue
                else:
                    # 可能是续行
                    pass
            else:
                table_rows.append(line)
                i += 1
                continue

        # 如果还在表格模式但遇到非表格行，先结束表格
        if in_table:
            add_table(doc, table_header, table_rows)
            in_table = False
            table_header = None
            table_rows = []

        # 标题
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            p = doc.add_heading('', level=level)
            add_rich_text(p, text)
            i += 1
            continue

        # 引用
        if line.strip().startswith('>'):
            text = re.sub(r'^>\s*', '', line.strip())
            p = doc.add_paragraph('', style='Quote')
            add_rich_text(p, text)
            i += 1
            continue

        # 列表项
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', line)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            text = list_match.group(3)
            p = doc.add_paragraph('')
            p.paragraph_format.left_indent = Cm(0.6 + indent_level * 0.6)
            # 添加符号
            marker = list_match.group(2)
            if marker in ['-', '*']:
                p.add_run('  ' * indent_level + '• ')
            else:
                p.add_run('  ' * indent_level + marker + ' ')
            add_rich_text(p, text)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph('')
        add_rich_text(p, line)
        i += 1

    # 处理文件末尾的表格
    if in_table:
        add_table(doc, table_header, table_rows)

    doc.save(docx_path)
    print(f"已生成: {docx_path}")


if __name__ == '__main__':
    base_dir = r'C:\Users\olina\Desktop\诉讼\庭审文档'
    files = [
        ('1-庭审策略文档（MECE重组版）v3.0.md', '1-庭审策略文档（MECE重组版）v3.0.docx'),
        ('发问提纲.md', '发问提纲.docx'),
    ]
    for md_name, docx_name in files:
        md_path = os.path.join(base_dir, md_name)
        docx_path = os.path.join(base_dir, docx_name)
        if os.path.exists(md_path):
            convert_md_to_docx(md_path, docx_path)
        else:
            print(f"文件不存在: {md_path}")
